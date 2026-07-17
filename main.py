# -*- coding: utf-8 -*-
"""
main.py - 招标信息聚合工具主入口
-------------------------------
用法:
  python main.py --input \"帮我找最近一周北京地区建筑类招标信息\" --mode cli
  python main.py --input \"每天早上8点帮我查上海IT类招标\" --mode cli
  python main.py --mode web  # 启动 Web UI
"""

import os
import sys
import time
import json
import logging
import argparse
import re
import subprocess
from datetime import datetime, timedelta
from pathlib import Path
from typing import List

# 将项目根目录加入路径
sys.path.insert(0, str(Path(__file__).parent))

import schedule
import threading
from concurrent.futures import ThreadPoolExecutor, TimeoutError

from intent_parser import IntentParser, TenderQuery
from docx_generator import generate_report, TenderRecord

# ---- 注册所有爬虫 ----
from scrapers.ccgp_scraper import CCGPScraper
from scrapers.ggzyjy_scraper import GGZYJYScraper
from scrapers.bidsearch_scraper import BidSearchScraper
from scrapers.search_engine_scraper import SearchEngineScraper

# 已注册的爬虫列表（agent 添加新爬虫后在这里注册即可）
SCRAPERS = [
    SearchEngineScraper(),  # 搜索引擎优先，最准
    CCGPScraper(),
    GGZYJYScraper(),
    BidSearchScraper(),
]

# ---- 日志 ----
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
    datefmt="%H:%M:%S",
)
logger = logging.getLogger("tender-agent")

OUTPUT_DIR = Path(__file__).parent / "output"


def ensure_output_dir():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


def run_once(user_input: str, query: TenderQuery = None, api_key: str = None) -> str:
    """
    执行一次完整的抓取→生成流程
    :param user_input: 用户的自然语言输入
    :param query: 已解析的查询参数（可选，不传则重新解析）
    :return: 生成的 Word 文件路径
    """
    _t0 = time.time()
    # 1. 意图解析
    if query is None:
        parser = IntentParser(api_key=api_key)
        logger.info(f"解析意图: {user_input}")
        query = parser.parse(user_input)
        logger.info(f"解析结果: {json.dumps(query.to_dict(), ensure_ascii=False)}")

    # 2. 并行抓取所有爬虫
    all_items = []
    with ThreadPoolExecutor(max_workers=min(4, len(SCRAPERS))) as pool:
        futures = {}
        for scraper in SCRAPERS:
            futures[pool.submit(
                scraper.fetch,
                region=query.region,
                industry=query.industry,
                time_range_days=query.time_range_days,
                raw_query=user_input,
            )] = scraper

        # 使用更宽松的超时，逐个等待
        import time as _time
        deadline = _time.time() + 90
        for future, scraper in [(future, futures[future]) for future in futures]:
            remaining = deadline - _time.time()
            if remaining <= 0:
                logger.warning(f"  {scraper.site_name} 总超时，跳过")
                future.cancel()
                continue
            try:
                items = future.result(timeout=min(remaining, 30))
                all_items.extend(items)
                logger.info(f"  {scraper.site_name} -> {len(items)} 条")
            except TimeoutError:
                logger.warning(f"  {scraper.site_name} 超时({remaining:.0f}s)，跳过")
            except Exception as e:
                logger.warning(f"  {scraper.site_name} 异常: {e}")

    # 3. 去重（按 source_url）
    seen_urls = set()
    deduped = []
    for item in all_items:
        if item.source_url not in seen_urls:
            seen_urls.add(item.source_url)
            deduped.append(item)

    logger.info(f"去重后共 {len(deduped)} 条 (原始 {len(all_items)} 条)")

    if not deduped:
        logger.info("真实网站未抓取到数据，使用演示数据展示流程")
        deduped = _demo_items(user_input, query)

    # 4. 生成 Word 报告
    ensure_output_dir()
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_path = OUTPUT_DIR / f"tender_report_{timestamp}.docx"

    records = [
        TenderRecord(
            title=item.title,
            region=item.region,
            industry=item.industry,
            amount=item.amount,
            publish_date=item.publish_date,
            deadline=item.deadline,
            source_url=item.source_url,
            source_site=item.source_site,
            summary=item.raw_text[:200] if item.raw_text else "",
        )
        for item in deduped
    ]

    elapsed = time.time() - _t0
    result_path = generate_report(records, user_input, str(output_path), elapsed_sec=elapsed)
    logger.info(f"报告已生成: {result_path}")
    return result_path



def _demo_items(user_input, query):
    from scrapers.base_scraper import RawTenderItem
    from datetime import datetime, timedelta
    import re
    t = datetime.now().strftime("%Y-%m-%d")
    d = (datetime.now() + timedelta(days=14)).strftime("%Y-%m-%d")
    r = query.region or ""
    ind = query.industry or ""

    # 提取核心关键词
    kw = user_input
    if r:
        kw = re.sub(rf"{r}[省市]?", "", kw)
    fillers = ["帮我","找找","查找","帮我找","查一下","一下","最近","看看",
               "搜索","一周","三天","今天","招标","信息","采购","公告",
               "地区","范围","类别的","什么","有没有","帮我查","请帮我",
               "每天","每日","每周","早上","上午","中午","下午","晚上",
               "点钟","点","分","立即","马上","现在","快速",
               "以内","以下","以上","及以上","之间","左右","大约",
               "建设项目","工程项目","项目","建设"]
    for w in sorted(fillers, key=lambda x: -len(x)):
        kw = kw.replace(w, "")
    kw = re.sub(r"\d+\s*[万亿千百]?\s*[元块]?", "", kw)
    kw = re.sub(r"[0-9一二三四五六七八九十]", "", kw)
    kw = re.sub(r"[的个之及与和类]", "", kw)
    kw = re.sub(r"\s+", "", kw)
    if not kw.strip():
        kw = "招标"
    else:
        kw = kw.strip()

    loc = r if r else "全国"

    note = ""
    if not ind:
        note = "（提示：当前查询未匹配到招标行业关键词，可尝试建筑/IT/医疗/市政等）"

    return [
        RawTenderItem(f"{loc} {kw} 相关项目A", r or "全国", ind or "综合", "约1200万元", t, d,
                      "(演示数据-非真实链接)", "【演示数据】",
                      f"此为演示数据。{note}".strip()),
        RawTenderItem(f"{loc} {kw} 配套采购B", r or "全国", ind or "综合", "约380万元", t, d,
                      "(演示数据-非真实链接)", "【演示数据】",
                      "系统支持多爬虫并行、robots.txt合规检查、自动去重。"),
        RawTenderItem(f"{loc} {kw} 建设项目C", r or "全国", ind or "综合", "约650万元", t, d,
                      "(演示数据-非真实链接)", "【演示数据】",
                      "实际使用时接入中国政府采购网、公共资源交易平台等真实网站。"),
    ]


def run_scheduled(user_input: str, api_key: str = None):
    """
    注册定时任务并保持运行
    """
    parser = IntentParser(api_key=api_key)
    query = parser.parse(user_input)

    if query.trigger_mode != "scheduled":
        logger.info("未检测到定时意图，执行单次抓取")
        return run_once(user_input, query)

    freq = query.schedule_freq or "daily"
    sched_time = query.schedule_time or "09:00"

    logger.info(f"注册定时任务: 每{freq} {sched_time} 执行")

    def job():
        logger.info(f"[定时任务触发] {datetime.now()}")
        try:
            path = run_once(user_input, query)
            logger.info(f"[定时任务完成] 报告: {path}")
        except Exception as e:
            logger.error(f"[定时任务失败] {e}")

    if freq == "daily":
        schedule.every().day.at(sched_time).do(job)
    elif freq == "weekly":
        schedule.every().monday.at(sched_time).do(job)
    else:  # once - 用 schedule 延迟到指定时间执行
        now = datetime.now()
        target = now.replace(
            hour=int(sched_time.split(":")[0]),
            minute=int(sched_time.split(":")[1]),
            second=0,
        )
        if target < now:
            target += timedelta(days=1)
        delay = (target - now).total_seconds()
        logger.info(f"单次定时任务将在 {target} 执行 (延迟 {delay:.0f} 秒)")
        schedule.every().day.at(sched_time).do(job)
        # 执行一次后取消
        schedule.every().day.at(sched_time).do(lambda: schedule.clear())

    logger.info("定时任务已就绪，等待触发... (按 Ctrl+C 停止)")
    try:
        while True:
            schedule.run_pending()
            time.sleep(30)
    except KeyboardInterrupt:
        logger.info("定时任务已停止")


def start_web_ui(host: str = "127.0.0.1", port: int = 5000):
    """启动 Web UI (Flask)"""
    try:
        from flask import Flask, render_template_string, request, send_file
    except ImportError:
        logger.error("Flask 未安装，请运行: pip install flask")
        sys.exit(1)

    app = Flask(__name__)

    HTML_TEMPLATE = """
<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>招标信息聚合工具</title>
    <style>
        * { margin: 0; padding: 0; box-sizing: border-box; }
        body {
            font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, sans-serif;
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            min-height: 100vh; display: flex; align-items: center; justify-content: center;
        }
        .container {
            background: white; border-radius: 16px; padding: 40px;
            max-width: 640px; width: 90%; box-shadow: 0 20px 60px rgba(0,0,0,0.3);
        }
        h1 { text-align: center; color: #333; margin-bottom: 8px; font-size: 24px; }
        .sub { text-align: center; color: #888; margin-bottom: 24px; font-size: 14px; }
        label { display: block; font-weight: 600; color: #555; margin-bottom: 8px; }
        input[type="text"] {
            width: 100%; padding: 14px 16px; border: 2px solid #e0e0e0;
            border-radius: 10px; font-size: 15px; transition: border-color 0.2s;
        }
        input[type="text"]:focus {
            border-color: #667eea; outline: none; box-shadow: 0 0 0 3px rgba(102,126,234,0.2);
        }
        .examples { margin-top: 10px; font-size: 12px; color: #999; }
        .examples span { cursor: pointer; color: #667eea; margin-right: 8px; text-decoration: underline; }
        button {
            width: 100%; padding: 14px; margin-top: 20px;
            background: linear-gradient(135deg, #667eea, #764ba2);
            color: white; border: none; border-radius: 10px;
            font-size: 16px; font-weight: 600; cursor: pointer; transition: transform 0.1s;
        }
        button:hover { transform: scale(1.02); }
        button:disabled { opacity: 0.6; cursor: not-allowed; transform: none; }
        .result { margin-top: 24px; padding: 16px; background: #f5f7fa; border-radius: 10px; display: none; }
        .result.show { display: block; }
        .result a { color: #667eea; font-weight: 600; font-size: 15px; }
        .error { color: #e74c3c; }
    </style>
</head>
<body>
    <div class="container">
        <h1>📋 招标信息聚合工具</h1>
        <p class="sub">自然语言驱动 · 多网站抓取 · Word 报告一键生成</p>

        <label for="query">描述你的需求</label>
        <input type="text" id="query" placeholder="例如：最近一周深圳市政工程招标">
        <div class="examples">
            试试：<span onclick="fill(this)">最近一周北京建筑类招标</span>
            <span onclick="fill(this)">深圳市政工程项目</span>
            <span onclick="fill(this)">全国IT类采购</span>
        </div>

        <button id="btn" onclick="generate()">🚀 生成报告</button>

        <div id="result" class="result"></div>

        <hr style="margin:24px 0 12px; border-color:#e0e0e0">
        <p class="sub">📅 定时监控（设置后每天自动搜，结果存 output/ 目录）</p>

        <label for="sched_query">监控查询词</label>
        <input type="text" id="sched_query" placeholder="例如：府谷县教育局 招标">

        <div style="display:flex; gap:10px; margin-top:12px">
            <div style="flex:1">
                <label for="sched_time">每天几点执行</label>
                <input type="time" id="sched_time" value="08:00" style="width:100%; padding:10px; border:2px solid #e0e0e0; border-radius:8px; font-size:14px">
            </div>
            <div style="flex:1">
                <label for="sched_freq">频率</label>
                <select id="sched_freq" style="width:100%; padding:10px; border:2px solid #e0e0e0; border-radius:8px; font-size:14px">
                    <option value="daily">每天</option>
                    <option value="weekly">每周一</option>
                </select>
            </div>
        </div>

        <div style="display:flex; gap:10px; margin-top:12px">
            <button id="sched_start" onclick="startSchedule()" style="background:linear-gradient(135deg,#27ae60,#2ecc71); flex:1">▶ 启动定时</button>
            <button id="sched_stop" onclick="stopSchedule()" style="background:linear-gradient(135deg,#e74c3c,#c0392b); flex:1; display:none">⏹ 停止</button>
        </div>
        <div style="text-align:center; margin-top:6px">
            <button onclick="saveToWindows()" style="background:#f0f0f0; border:1px solid #ccc; border-radius:8px; padding:8px 16px; font-size:13px; cursor:pointer">💾 保存到Windows定时任务 (关闭浏览器也能跑)</button>
        </div>
        <div id="sched_status" style="margin-top:8px; font-size:13px; color:#888; text-align:center"></div>
    </div>

    <script>
        function fill(el) { document.getElementById('query').value = el.textContent; }
        async function generate() {
            const query = document.getElementById('query').value.trim();
            if (!query) return;

            const btn = document.getElementById('btn');
            const result = document.getElementById('result');
            btn.disabled = true;
            btn.textContent = '正在抓取并生成报告...';
            result.className = 'result';
            result.textContent = '';

            try {
                const resp = await fetch('/api/generate', {
                    method: 'POST',
                    headers: {'Content-Type': 'application/json'},
                    body: JSON.stringify({query: query})
                });
                const data = await resp.json();
                result.className = 'result show';
                if (data.success) {
                    result.innerHTML = '[OK] 报告生成成功！<br><a href=\"' + data.download + '\">📥 下载报告</a><br><small>共 ' + data.count + ' 条信息</small>';
                } else {
                    result.innerHTML = '<span class=\"error\">❌ ' + data.error + '</span>';
                }
            } catch(e) {
                result.className = 'result show';
                result.innerHTML = '<span class=\"error\">❌ 请求失败: ' + e.message + '</span>';
            } finally {
                btn.disabled = false;
                btn.textContent = '🚀 生成报告';
            }
        }
        async function startSchedule() {
            const q = document.getElementById('sched_query').value.trim();
            const t = document.getElementById('sched_time').value;
            const f = document.getElementById('sched_freq').value;
            if (!q) return alert('请输入监控查询词');

            const btn = document.getElementById('sched_start');
            btn.disabled = true;
            btn.textContent = '启动中...';

            try {
                const resp = await fetch('/api/schedule/start', {
                    method: 'POST',
                    headers: {'Content-Type': 'application/json'},
                    body: JSON.stringify({query: q, time: t, freq: f})
                });
                const data = await resp.json();
                if (data.success) {
                    document.getElementById('sched_start').style.display = 'none';
                    document.getElementById('sched_stop').style.display = 'block';
                    document.getElementById('sched_status').textContent = '✅ 已启动: 每' + (f==='daily'?'天':'周一') + ' ' + t + ' 自动搜「' + q + '」';
                }
            } catch(e) {
                document.getElementById('sched_status').textContent = '❌ 失败: ' + e.message;
            }
            btn.disabled = false;
        }
        async function stopSchedule() {
            try {
                const resp = await fetch('/api/schedule/stop', {method: 'POST'});
                const data = await resp.json();
                if (data.success) {
                    document.getElementById('sched_start').style.display = 'block';
                    document.getElementById('sched_stop').style.display = 'none';
                    document.getElementById('sched_status').textContent = '⏹ 已停止';
                }
            } catch(e) {}
        }
        async function saveToWindows() {
            const q = document.getElementById('sched_query').value.trim();
            const t = document.getElementById('sched_time').value;
            const f = document.getElementById('sched_freq').value;
            const st = document.getElementById('sched_status');
            if (!q) { st.textContent = '请先输入监控查询词'; return; }

            st.textContent = '正在保存...';
            try {
                const resp = await fetch('/api/schedule/save', {
                    method: 'POST',
                    headers: {'Content-Type': 'application/json'},
                    body: JSON.stringify({query: q, time: t, freq: f})
                });
                const data = await resp.json();
                if (data.success) {
                    st.innerHTML = '✅ 已保存！<br>任务: ' + data.name + '<br>每天 ' + t + ' 自动执行<br><small>关闭浏览器不受影响 | 查看: Win+R → taskschd.msc</small>';
                } else {
                    st.innerHTML = '❌ 失败: ' + data.error;
                }
            } catch(e) {
                st.innerHTML = '❌ 网络错误: ' + e.message;
            }
        }
    </script>
</body>
</html>
"""

    @app.route("/")
    def index():
        return render_template_string(HTML_TEMPLATE)

    @app.route("/api/generate", methods=["POST"])
    def api_generate():
        data = request.get_json()
        user_input = data.get("query", "").strip()
        # Security: read API key from environment variable
        api_key = os.getenv("ANTHROPIC_API_KEY")
        if not user_input:
            return {"success": False, "error": "请输入查询内容"}

        try:
            result_path = run_once(user_input, api_key=api_key)
            filename = os.path.basename(result_path)
            # 统计记录数
            from docx import Document
            doc = Document(result_path)
            table = doc.tables[0] if doc.tables else None
            count = len(table.rows) - 1 if table else 0
            return {
                "success": True,
                "download": f"/download/{filename}",
                "count": count,
            }
        except Exception as e:
            logger.exception("生成失败")
            return {"success": False, "error": str(e)}

    @app.route("/download/<filename>")
    def download(filename):
        return send_file(
            OUTPUT_DIR / filename,
            as_attachment=True,
            download_name=filename,
        )

    # ---- 定时任务管理 ----
    scheduled_jobs = {}
    jobs_lock = threading.Lock()

    @app.route("/api/schedule/start", methods=["POST"])
    def api_schedule_start():
        data = request.get_json()
        q = data.get("query", "").strip()
        t = data.get("time", "08:00")
        freq = data.get("freq", "daily")
        if not q:
            return {"success": False, "error": "请输入查询词"}

        # 清除旧任务
        schedule.clear()
        with jobs_lock:
            scheduled_jobs.clear()

        def job():
            logger.info(f"[定时任务] {datetime.now()}: {q}")
            try:
                path = run_once(q)
                logger.info(f"[定时任务] 报告: {path}")
            except Exception as e:
                logger.error(f"[定时任务] 失败: {e}")

        if freq == "daily":
            schedule.every().day.at(t).do(job)
        else:
            schedule.every().monday.at(t).do(job)

        with jobs_lock:
            scheduled_jobs["active"] = True

        # 在后台线程跑 schedule
        def run_schedule_loop():
            while True:
                with jobs_lock:
                    if not scheduled_jobs.get('active'):
                        break
                schedule.run_pending()
                time.sleep(30)
        sched_thread = threading.Thread(target=run_schedule_loop, daemon=True)
        sched_thread.start()

        return {"success": True, "msg": f"已启动: 每{freq} {t} 搜「{q}」"}

    @app.route("/api/schedule/save", methods=["POST"])
    def api_schedule_save():
        data = request.get_json()
        q = data.get("query", "").strip()
        t = data.get("time", "08:00")
        freq = data.get("freq", "daily")
        if not q:
            return {"success": False, "error": "请输入查询词"}

        safe_q = re.sub(r"[^\w一-鿿\s]", "", q)[:20]
        task_name = f"TenderAgent_{safe_q.replace(' ','_')}"
        python_path = sys.executable
        script_path = os.path.join(os.path.dirname(__file__), "main.py")
        h, m = t.split(":")

        # 生成 .bat 脚本
        bat_safe_q = q.replace('"', '""')  # Windows batch escape
        bat_content = f'''@echo off
chcp 65001 >nul
cd /d "{os.path.dirname(__file__)}"
echo [%date% %time%] TenderAgent: {bat_safe_q}
"{python_path}" "{script_path}" -i "{bat_safe_q}" -m cli
echo Done. Report saved to output\
'''
        bat_path = os.path.join(OUTPUT_DIR, f"scheduled_{task_name}.bat")
        with open(bat_path, "w", encoding="utf-8") as bf:
            bf.write(bat_content)

        # 尝试创建 Windows 定时任务
        exe_cmd = f'cmd /c "{bat_path}"'
        sc_param = 'WEEKLY' if freq == 'weekly' else 'DAILY'
        cmd = (
            f'schtasks /Create /SC {sc_param} /ST {h}:{m} /TN "{task_name}" '
            f'/TR "{exe_cmd}" /F'
        )
        try:
            result = subprocess.run(cmd, shell=True, capture_output=True, text=True, timeout=15)
            if result.returncode == 0:
                msg = f"已创建Windows定时任务: 每天 {t} 自动执行。关闭浏览器不受影响。"
                return {"success": True, "name": task_name, "msg": msg}
        except Exception as e:
            logger.warning(f'schtasks failed: {e}')

        # schtasks 失败：退回手动方案
        msg = (f"定时任务脚本已生成。请右键点击下方文件→以管理员身份运行 来注册:\n"
               f"{bat_path}\n\n"
               f"或者手动添加Windows定时任务:\n"
               f"Win+R → taskschd.msc → 创建任务 → 触发器:每天{t} → 操作:启动程序 \"{bat_path}\"")
        return {"success": True, "name": task_name, "msg": msg, "bat_path": bat_path}

    @app.route("/api/schedule/stop", methods=["POST"])
    def api_schedule_stop():
        schedule.clear()
        with jobs_lock:
            scheduled_jobs.clear()
        return {"success": True, "msg": "已停止"}

    logger.info(f"Web UI 启动: http://{host}:{port}")
    app.run(host=host, port=port, debug=False)


def main():
    # Windows 终端 UTF-8 支持
    import io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
    parser = argparse.ArgumentParser(
        description="招标信息聚合工具 - 自然语言驱动，多网站抓取，自动生成Word报告"
    )
    parser.add_argument(
        "--input", "-i", type=str,
        help="自然语言查询输入，例如: '帮我找最近一周北京地区建筑类招标信息'",
    )
    parser.add_argument(
        "--mode", "-m", type=str, default="cli",
        choices=["cli", "scheduled", "web"],
        help="运行模式: cli(单次), scheduled(定时), web(Web界面)",
    )
    parser.add_argument(
        "--port", "-p", type=int, default=5000,
        help="Web UI 端口 (默认 5000)",
    )
    parser.add_argument(
        "--all", action="store_true",
        help="抓取全部爬虫，不做跳过",
    )

    args = parser.parse_args()

    # 模式分发
    if args.mode == "web":
        start_web_ui(port=args.port)
    elif args.mode == "scheduled":
        if not args.input:
            print("错误: scheduled 模式需要 --input 参数")
            sys.exit(1)
        run_scheduled(args.input)
    else:  # cli
        if not args.input:
            print("错误: CLI 模式需要 --input 参数")
            print('示例: python main.py --input "帮我找最近一周北京地区建筑类招标信息"')
            sys.exit(1)
        result_path = run_once(args.input)
        print(f"\n[OK] 报告已生成: {result_path}")


if __name__ == "__main__":
    main()

