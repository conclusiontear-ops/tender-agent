# -*- coding: utf-8 -*-
"""
Word 文档生成模块 V2 - 企业报告模板
"""
from collections import Counter, defaultdict
from datetime import datetime
from typing import List, Dict, Optional
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml.shared import OxmlElement, qn

# ---- 品牌色系统 ----
PRIMARY  = RGBColor(0x25, 0x63, 0xEB)   # #2563EB 主蓝
DARK     = RGBColor(0x1E, 0x29, 0x3B)   # 正文黑
MUTED    = RGBColor(0x6B, 0x72, 0x80)   # 辅助灰
ACCENT   = RGBColor(0x05, 0x63, 0xC1)   # 超链接蓝
LIGHT_BG = "EFF6FF"                      # 浅蓝底（用于AI分析框）


class TenderRecord:
    def __init__(self, title, region, industry, amount, publish_date, deadline,
                 source_url, source_site, summary=""):
        self.title = title
        self.region = region
        self.industry = industry
        self.amount = amount
        self.publish_date = publish_date
        self.deadline = deadline
        self.source_url = source_url
        self.source_site = source_site
        self.summary = summary


# ─── 工具函数 ───────────────────────────────────────

def _hyperlink(paragraph, url, text, color=ACCENT, size=Pt(9)):
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hl = OxmlElement("w:hyperlink")
    hl.set(qn("r:id"), r_id)
    nr = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    for tag, val in [("w:color", color), ("w:u", "single"),
                      ("w:sz", str(int(size.pt * 2))), ("w:rFonts", "Microsoft YaHei")]:
        el = OxmlElement(tag)
        if tag == "w:rFonts":
            el.set(qn("w:ascii"), "Microsoft YaHei")
            el.set(qn("w:hAnsi"), "Microsoft YaHei")
        else:
            el.set(qn("w:val"), str(val))
        rPr.append(el)
    nr.append(rPr)
    nr.text = text
    hl.append(nr)
    paragraph._p.append(hl)
    return paragraph


def _set_cell_shading(cell, color):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def _set_paragraph_shading(paragraph, color):
    pPr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    shading.set(qn("w:val"), "clear")
    pPr.append(shading)


def _hdr_footer(doc):
    """页眉 + 页脚 (Page X / Y)"""
    for section in doc.sections:
        # 页眉
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = hp.add_run("Tender Agent  ·  智能招投标分析报告")
        r.font.size = Pt(8); r.font.color.rgb = MUTED

        # 页脚 - Page X / Y
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = fp.add_run("Page ")
        r1.font.size = Pt(8); r1.font.color.rgb = MUTED
        _add_field(fp, "PAGE")
        r2 = fp.add_run(" / ")
        r2.font.size = Pt(8); r2.font.color.rgb = MUTED
        _add_field(fp, "NUMPAGES")
        r3 = fp.add_run("  |  AI Future Talent Competition")
        r3.font.size = Pt(8); r3.font.color.rgb = MUTED


def _add_field(paragraph, field_name):
    run = OxmlElement("w:r")
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    run.append(fldChar1)
    instrText = OxmlElement("w:instrText")
    instrText.text = f" {field_name} "
    run.append(instrText)
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run.append(fldChar2)
    rPr = OxmlElement("w:rPr")
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "16")
    rPr.append(sz)
    run.append(rPr)
    paragraph._p.append(run)


def _confidential_stamp(doc):
    """TOC 页水印标记"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("CONFIDENTIAL  ·  Internal Use Only")
    r.font.size = Pt(8); r.font.color.rgb = MUTED; r.font.italic = True
    doc.add_paragraph()


# ─── 各模块 ─────────────────────────────────────────

def _page_cover(doc, analysis):
    """封面"""
    for _ in range(5): doc.add_paragraph()

    h = doc.add_paragraph(); h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = h.add_run("招投标信息分析报告")
    r.font.size = Pt(30); r.font.bold = True; r.font.color.rgb = PRIMARY

    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Tender Agent  ·  AI-Powered Insights")
    r.font.size = Pt(14); r.font.color.rgb = MUTED

    doc.add_paragraph()

    for label, value in [
        ("Report ID", analysis.get("report_id", "")),
        ("查询条件", analysis.get("query_desc", "")),
        ("生成时间", analysis.get("generated_at", "")),
        ("检索结果", f"共 {analysis.get('total', 0)} 条"),
    ]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(f"{label}："); r1.font.size = Pt(11); r1.font.color.rgb = MUTED
        r2 = p.add_run(value); r2.font.size = Pt(11); r2.font.color.rgb = DARK

    for _ in range(5): doc.add_paragraph()
    bot = doc.add_paragraph(); bot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = bot.add_run("AI Future Talent Competition"); r.font.size = Pt(10); r.font.color.rgb = MUTED

    doc.add_page_break()


def _page_toc(doc, analysis):
    """目录 + Confidential"""
    _confidential_stamp(doc)
    doc.add_heading("目  录", level=1)
    toc_items = [
        "1. 查询摘要",
        "2. AI 综合分析",
        "3. 数据统计",
        "4. 招标详情",
        "5. AI 建议",
    ]
    for item in toc_items:
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item); r.font.size = Pt(12); r.font.color.rgb = PRIMARY
    doc.add_page_break()


def _page_summary(doc, analysis):
    """查询摘要"""
    doc.add_heading("查询摘要", level=1)

    items = [
        ("查询条件", analysis.get("query_desc", "")),
        ("Report ID", analysis.get("report_id", "")),
        ("生成时间", analysis.get("generated_at", "")),
        ("检索结果", f"共获取 {analysis.get('total', 0)} 条有效信息"),
        ("数据来源", f"{len(analysis.get('sources', {}))} 个平台"),
    ]
    for label, value in items:
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run(f"{label}："); r1.font.bold = True; r1.font.size = Pt(10.5); r1.font.color.rgb = DARK
        r2 = p.add_run(value); r2.font.size = Pt(10.5); r2.font.color.rgb = MUTED

    # 来源详情
    sources = analysis.get("sources", {})
    if sources:
        doc.add_paragraph()
        p = doc.add_paragraph()
        r = p.add_run("访问网站：")
        r.font.bold = True; r.font.size = Pt(10)
        for src, cnt in sources.items():
            p2 = doc.add_paragraph()
            r = p2.add_run(f"  ✓ {src}  —  {cnt} 条")
            r.font.size = Pt(9.5); r.font.color.rgb = DARK

    # 耗时/状态
    elapsed = analysis.get("elapsed_sec", 0)
    doc.add_paragraph()
    p3 = doc.add_paragraph()
    r = p3.add_run(f"查询耗时：{elapsed:.1f} 秒    状态：✓ 完成")
    r.font.size = Pt(9.5); r.font.color.rgb = MUTED

    doc.add_page_break()


def _page_ai_analysis(doc, analysis):
    """AI 综合分析"""
    doc.add_heading("AI 综合分析", level=1)

    ai_text = analysis.get("ai_summary", "暂无分析数据。")
    
    # 浅蓝底色框
    p = doc.add_paragraph()
    _set_paragraph_shading(p, LIGHT_BG)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.right_indent = Cm(1)

    for paragraph_text in ai_text.split("\n\n"):
        if paragraph_text.strip():
            r = p.add_run(paragraph_text.strip() + "\n")
            r.font.size = Pt(10.5); r.font.color.rgb = DARK

    doc.add_paragraph()
    
    # AI 说明
    note = doc.add_paragraph()
    r = note.add_run("※ 以上分析由 AI 自动生成，仅供参考，不构成投标决策建议。")
    r.font.size = Pt(8); r.font.italic = True; r.font.color.rgb = MUTED

    doc.add_page_break()


def _page_statistics(doc, analysis):
    """数据统计"""
    doc.add_heading("数据统计", level=1)

    stats_config = [
        ("来源统计", analysis.get("sources", {})),
        ("地区统计", analysis.get("regions", {})),
        ("行业统计", analysis.get("industries", {})),
    ]
    
    for title, data in stats_config:
        if not data:
            continue
        doc.add_heading(title, level=3)
        table = doc.add_table(rows=1, cols=2)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "类别"; hdr[1].text = "数量"
        for cell in hdr:
            _set_cell_shading(cell, "2563EB")
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.bold = True; r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    r.font.size = Pt(9)
        for k, v in sorted(data.items(), key=lambda x: -x[1]):
            row = table.add_row().cells
            row[0].text = k; row[1].text = str(v)
            for cell in row:
                for p in cell.paragraphs:
                    for r in p.runs: r.font.size = Pt(9)
        doc.add_paragraph()

    # 金额分层
    buckets = analysis.get("amount_buckets", {})
    if buckets:
        doc.add_heading("预算分布", level=3)
        table2 = doc.add_table(rows=1, cols=2)
        table2.style = "Light Grid Accent 1"
        h2 = table2.rows[0].cells
        h2[0].text = "预算区间"; h2[1].text = "数量"
        for cell in h2:
            _set_cell_shading(cell, "2563EB")
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.bold = True; r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    r.font.size = Pt(9)
        for k, v in buckets.items():
            if v > 0:
                row = table2.add_row().cells
                row[0].text = k; row[1].text = str(v)
                for cell in row:
                    for p in cell.paragraphs:
                        for r in p.runs: r.font.size = Pt(9)

    doc.add_page_break()


def _page_projects(doc, records):
    """招标详情 - 项目卡片风格"""
    doc.add_heading("招标详情", level=1)

    by_source = defaultdict(list)
    for rec in records:
        by_source[rec.source_site].append(rec)

    project_num = 0
    for src_name, src_records in by_source.items():
        doc.add_heading(f"来源：{src_name}（{len(src_records)} 条）", level=2)

        for rec in src_records:
            project_num += 1

            # 分隔线
            divider = doc.add_paragraph()
            r = divider.add_run("─" * 50)
            r.font.size = Pt(6); r.font.color.rgb = MUTED

            # 项目编号
            p_num = doc.add_paragraph()
            r = p_num.add_run(f"项目 {project_num}")
            r.font.bold = True; r.font.size = Pt(11); r.font.color.rgb = PRIMARY

            # 项目名称
            p_title = doc.add_paragraph()
            r = p_title.add_run(rec.title)
            r.font.bold = True; r.font.size = Pt(12); r.font.color.rgb = DARK

            # 信息字段
            fields = [
                ("地区", rec.region or "全国"),
                ("行业", rec.industry or "不限"),
                ("预算金额", rec.amount or "详见公告"),
                ("发布时间", rec.publish_date),
                ("截止日期", rec.deadline),
                ("来源", rec.source_site),
            ]
            for label, value in fields:
                p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(1)
                r1 = p.add_run(f"  {label}："); r1.font.bold = True; r1.font.size = Pt(9.5); r1.font.color.rgb = DARK
                r2 = p.add_run(value); r2.font.size = Pt(9.5); r2.font.color.rgb = MUTED

            # 链接
            if rec.source_url and rec.source_url.startswith("http"):
                p_link = doc.add_paragraph()
                _hyperlink(p_link, rec.source_url, f"  🔗 原文链接", size=Pt(9))

            # 摘要
            if rec.summary:
                p_sum = doc.add_paragraph()
                r = p_sum.add_run(f"  AI 摘要：{rec.summary[:200]}")
                r.font.size = Pt(9); r.font.italic = True; r.font.color.rgb = MUTED

            # 可信度（基于链接可访问性）
            trust = "★★★★★" if rec.source_url and rec.source_url.startswith("http") else "★★★☆☆"
            p_trust = doc.add_paragraph()
            r = p_trust.add_run(f"  可信度：{trust}")
            r.font.size = Pt(8); r.font.color.rgb = MUTED

            doc.add_paragraph()

    doc.add_page_break()


def _page_suggestions(doc, analysis):
    """AI 建议"""
    doc.add_heading("AI 建议", level=1)

    suggestions = analysis.get("ai_suggestions", [])
    if not suggestions:
        suggestions = ["暂无建议。请检查查询条件后重试。"]

    for s in suggestions:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(f"✓ {s}")
        r.font.size = Pt(10.5); r.font.color.rgb = DARK

    doc.add_paragraph()
    note = doc.add_paragraph()
    r = note.add_run("※ 以上建议由 AI 自动生成，请结合实际情况进行判断。")
    r.font.size = Pt(8); r.font.italic = True; r.font.color.rgb = MUTED


# ─── 主入口 ─────────────────────────────────────────

def generate_report(records: List[TenderRecord], query_desc: str, output_path: str,
                    analysis: Optional[Dict] = None, elapsed_sec: float = 0) -> str:
    if analysis is None:
        from report_analyzer import analyze
        analysis = analyze(records, query_desc, elapsed_sec)

    doc = Document()

    # 全局样式
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(10.5)
    style.font.color.rgb = DARK

    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    _hdr_footer(doc)

    if not records:
        _page_cover(doc, analysis)
        doc.add_paragraph("本次未检索到符合条件的招标信息，请调整查询条件后重试。")
        doc.save(output_path)
        return output_path

    # 按顺序生成各页
    _page_cover(doc, analysis)
    _page_toc(doc, analysis)
    _page_summary(doc, analysis)
    _page_ai_analysis(doc, analysis)
    _page_statistics(doc, analysis)
    _page_projects(doc, records)
    _page_suggestions(doc, analysis)

    doc.save(output_path)
    return output_path
